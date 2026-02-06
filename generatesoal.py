import os
from google import genai
from dotenv import load_dotenv
load_dotenv()
import re
import json
import base64
import requests
import mysql.connector
from flask import Flask, render_template, request, redirect, url_for, session, send_file
import PyPDF2
import docx
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
app.secret_key = "rahasia123"

# ---------------- MYSQL CONNECTION ----------------
db = mysql.connector.connect(
    host="localhost",
    user="root",          # ganti sesuai user mysql kamu
    password="",          # password MySQL
    database="sekolah"    # nama database
)
cursor = db.cursor(dictionary=True)

# ---------------- AUTH ----------------

@app.route("/")
def home():
    return redirect(url_for ("login"))
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        cursor.execute("SELECT * FROM users WHERE username=%s AND password=%s", (username, password))
        user = cursor.fetchone()

        if user:
            session["user"] = user
            if user["role"] == "pengajar":
                return redirect(url_for("dashboard_pengajar"))
            elif user["role"] == "murid":
                return redirect(url_for("dashboard_murid"))
        else:
            return render_template("login2.html", error="Username atau password salah!")

    return render_template("login2.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        role = request.form["role"]

        try:
            cursor.execute(
                "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
                (username, password, role)
            )
            db.commit()
            # arahkan ke halaman register lagi dengan status success
            return redirect(url_for("register", status="success"))
        except Exception as e:
            print("Error:", e)
            return redirect(url_for("register", status="error"))

    return render_template("register2.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------------- DASHBOARD ----------------
@app.route("/dashboard_pengajar")
def dashboard_pengajar():
    if "user" not in session or session["user"]["role"] != "pengajar":
        return redirect(url_for("login"))
    return render_template("DashboardPengajar.html")

@app.route("/dashboard_murid")
def dashboard_murid():
    if "user" not in session or session["user"]["role"] != "murid":
        return redirect(url_for("login"))

    murid_id = session["user"]["id"]

    # Ujian yang sudah dikerjakan
    cursor.execute("""
        SELECT u.judul, u.mapel, h.status, h.nilai
        FROM ujian u
        JOIN hasil_ujian h ON u.id = h.ujian_id
        WHERE h.murid_id = %s AND h.status = 'selesai'
    """, (murid_id,))
    ujian_selesai = cursor.fetchall()

    # Ujian yang belum dikerjakan
    cursor.execute("""
        SELECT u.id, u.judul, u.mapel, u.jumlah_soal
        FROM ujian u
        LEFT JOIN hasil_ujian h 
        ON u.id = h.ujian_id AND h.murid_id = %s
        WHERE h.id IS NULL OR h.status = 'belum'
    """, (murid_id,))
    ujian_belum = cursor.fetchall()

    return render_template("dashboard_murid.html",
                           ujian_selesai=ujian_selesai,
                           ujian_belum=ujian_belum)


@app.route("/ujian/<int:ujian_id>", methods=["GET", "POST"])
def kerjakan_ujian(ujian_id):
    if "user" not in session or session["user"]["role"] != "murid":
        return redirect(url_for("login"))

    murid_id = session["user"]["id"]

    # --- TAMBAHKAN LINE INI UNTUK MEMPERBARUI KONEKSI DATABASE ---
    db.commit()  # Refresh koneksi agar bisa baca data terbaru
    cursor.execute("SELECT * FROM ujian WHERE id=%s", (ujian_id,))
    ujian = cursor.fetchone()

    # Ambil soal
    cursor.execute("SELECT * FROM soal WHERE ujian_id=%s ORDER BY id ASC", (ujian_id,))
    soal_list = cursor.fetchall()

    # --- PASANG CCTV (DEBUGGING) DI SINI ---
    print("\n" + "="*30)
    print(f"DEBUG: ID Ujian dari URL = {ujian_id}")
    if soal_list:
        print(f"DEBUG: Ditemukan {len(soal_list)} soal.")
        print(f"DEBUG: Contoh Soal Pertama: {soal_list[0]}") 
        # Lihat nama key di dictionary ini (apakah 'soal', 'pertanyaan', atau lainnya?)
    else:
        print("DEBUG: ❌ Soal KOSONG/Tidak Ditemukan untuk ID ini!")
    print("="*30 + "\n")
    # ---------------------------------------

    if request.method == "POST":
        hasil_per_soal = []
        benar_count = 0

        for i, soal in enumerate(soal_list, start=1):
            jawaban_murid = request.form.get(f"jawaban{i}", "").strip().lower()
            jawaban_benar = soal["jawaban"].strip().lower()

            # cek cosine similarity
            skor = cek_kemiripan(jawaban_murid, jawaban_benar)

            if skor >= 0.5:  # dianggap benar
                status = f"✅ Benar (similarity {skor:.2f})"
                feedback = "Jawabanmu sudah sesuai dengan kunci jawaban."
                benar_count += 1
            else:  # salah → kasih feedback
                status = f"❌ Salah (similarity {skor:.2f})"

                # bikin feedback detail
                if not jawaban_murid:
                    feedback = "Kamu belum menjawab soal ini."
                else:
                    feedback = (
                        f"Jawabanmu kurang tepat. "
                        f"Bagian penting yang seharusnya ada: '{jawaban_benar}'. "
                        f"Coba perhatikan pembahasan untuk memahami lebih baik."
                    )

            hasil_per_soal.append({
                "no": i,
                "soal": soal["soal"],
                "jawaban_murid": jawaban_murid,
                "jawaban_benar": soal["jawaban"],
                "pembahasan": soal["pembahasan"],
                "status": status,
                "feedback": feedback
            })

        # Hitung nilai (100 dibagi jumlah soal)
        nilai = int((benar_count / len(soal_list)) * 100)

        # Simpan hasil ke DB
        cursor.execute("""
            INSERT INTO hasil_ujian (ujian_id, murid_id, nilai, status)
            VALUES (%s, %s, %s, 'selesai')
            ON DUPLICATE KEY UPDATE nilai=%s, status='selesai'
        """, (ujian_id, murid_id, nilai, nilai))
        db.commit()

        return render_template("hasil_ujian.html",
                               ujian=ujian,
                               hasil=hasil_per_soal,
                               nilai=nilai,
                               total=len(soal_list))

    # ... (lanjutan kode POST kamu) ...
    
    return render_template("kerjakan_ujian.html", ujian=ujian, soal_list=soal_list)


# @app.route("/ujian/<int:ujian_id>", methods=["GET", "POST"])
# def kerjakan_ujian(ujian_id):
#     if "user" not in session or session["user"]["role"] != "murid":
#         return redirect(url_for("login"))

#     murid_id = session["user"]["id"]

#     # Data ujian
#     cursor.execute("SELECT * FROM ujian WHERE id=%s", (ujian_id,))
#     ujian = cursor.fetchone()

#     # Ambil soal dari tabel soal
#     cursor.execute("SELECT * FROM soal WHERE ujian_id=%s ORDER BY id ASC", (ujian_id,))
#     soal_list = cursor.fetchall()

#     if request.method == "POST":
#         hasil_per_soal = []
#         benar_count = 0

#         for i, soal in enumerate(soal_list, start=1):
#             jawaban_murid = request.form.get(f"jawaban{i}", "").strip().lower()
#             jawaban_benar = soal["jawaban"].strip().lower()

#             # cek cosine similarity
#             skor = cek_kemiripan(jawaban_murid, jawaban_benar)

#             if skor >= 0.5:  # dianggap benar
#                 status = f"✅ Benar (similarity {skor:.2f})"
#                 feedback = "Jawabanmu sudah sesuai dengan kunci jawaban."
#                 benar_count += 1
#             else:  # salah → kasih feedback
#                 status = f"❌ Salah (similarity {skor:.2f})"

#                 # bikin feedback detail
#                 if not jawaban_murid:
#                     feedback = "Kamu belum menjawab soal ini."
#                 else:
#                     feedback = (
#                         f"Jawabanmu kurang tepat. "
#                         f"Bagian penting yang seharusnya ada: '{jawaban_benar}'. "
#                         f"Coba perhatikan pembahasan untuk memahami lebih baik."
#                     )

#             hasil_per_soal.append({
#                 "no": i,
#                 "soal": soal["soal"],
#                 "jawaban_murid": jawaban_murid,
#                 "jawaban_benar": soal["jawaban"],
#                 "pembahasan": soal["pembahasan"],
#                 "status": status,
#                 "feedback": feedback
#             })

#         # Hitung nilai (100 dibagi jumlah soal)
#         nilai = int((benar_count / len(soal_list)) * 100)

#         # Simpan hasil ke DB
#         cursor.execute("""
#             INSERT INTO hasil_ujian (ujian_id, murid_id, nilai, status)
#             VALUES (%s, %s, %s, 'selesai')
#             ON DUPLICATE KEY UPDATE nilai=%s, status='selesai'
#         """, (ujian_id, murid_id, nilai, nilai))
#         db.commit()

#         return render_template("hasil_ujian.html",
#                                ujian=ujian,
#                                hasil=hasil_per_soal,
#                                nilai=nilai,
#                                total=len(soal_list))

#     # kalau belum submit → tampilkan soal
#     return render_template("kerjakan_ujian.html", ujian=ujian, soal_list=soal_list)



# ---------------- FITUR BUAT SOAL ----------------
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_word(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def call_gemini_api(prompt: str) -> str:
    try:
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=prompt
        )
        return response.text
    except Exception as e:
        return f"⚠ Error Gemini: {str(e)}"


def cek_kemiripan(jawaban_murid, jawaban_benar):
    if not jawaban_murid.strip():
        return 0.0
    vectorizer = TfidfVectorizer().fit([jawaban_murid, jawaban_benar])
    tfidf_matrix = vectorizer.transform([jawaban_murid, jawaban_benar])
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
    return similarity[0][0]  # nilai antara 0 - 1



@app.route("/upload_modul", methods=["GET", "POST"])
def upload_modul():
    if "user" not in session or session["user"]["role"] != "pengajar":
        return redirect(url_for("login"))

    soal_output = None
    if request.method == "POST":
        uploaded_file = request.files["file"]
        kelas = request.form.get("kelas")
        mapel = request.form.get("mapel")
        semester = request.form.get("semester")
        judul = request.form.get("judul")
        jumlah_soal = int(request.form.get("jumlah_soal", 5))

        # Ambil teks dari file
        # Ambil teks dari file
        if uploaded_file.filename.endswith(".pdf"):
            materi = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.filename.endswith(".docx"):
            materi = extract_text_from_word(uploaded_file)
        else:
            materi = None

        # HEMAT TOKEN (WAJIB)
        if materi:
            materi = materi[:6000]


        tingkat_kesulitan = {
            "X": "mudah hingga menengah, sesuai pemahaman dasar.",
            "XI": "menengah, sesuai dengan siswa tingkat lanjut menengah.",
            "XII": "menengah hingga sulit, sesuai persiapan ujian akhir."
        }

        
        if materi:
                # Tambahkan instruksi agar rumus matematika ditandai sebagai LaTeX
                prompt = f"""
                Kamu adalah AI pembuat soal esai resmi untuk Ujian Nasional.
                Konteks (materi modul/buku): {materi}
                Instruksi:
                1. Buatkan {jumlah_soal} soal berdasarkan materi di atas.
                2. Tingkat kesulitan soal harus {tingkat_kesulitan.get(kelas, 'menengah')}.
                3. Untuk setiap soal, sertakan jawaban dan pembahasan.
                4. Jika soal memuat rumus matematika, tuliskan rumus menggunakan LaTeX dan bungkus dengan $...$ untuk inline atau $$...$$ untuk display.
                5. Gunakan format tanpa simbol Markdown.
                6. Format wajib:
                   Soal 1: ...
                   Jawaban: ...
                   Pembahasan: ...
                Gunakan bahasa Indonesia baku.
                """
                soal_output = call_gemini_api(prompt)

                tipe_soal = request.form.get("tipe_soal") # Tambahkan input ini di HTML form

    if materi:
        materi = materi[:6000]
        
        # Modifikasi Prompt
        if tipe_soal == "pilihan_ganda":
            format_instruksi = """
            6. Format wajib:
               Soal 1: ...
               A: ...
               B: ...
               C: ...
               D: ...
               Jawaban: (Hanya huruf A/B/C/D)
               Pembahasan: ...
            """
        else:
            format_instruksi = """
            6. Format wajib:
               Soal 1: ...
               Jawaban: ...
               Pembahasan: ...
            """

        prompt = f"""
        Kamu adalah AI pembuat soal {tipe_soal} resmi.
        Materi: {materi}
        Buatkan {jumlah_soal} soal {tipe_soal}.
        {format_instruksi}
        Gunakan bahasa Indonesia baku.
        """
        soal_output = call_gemini_api(prompt)

                # Jangan langsung menyimpan ke DB — tampilkan preview terlebih dahulu
        return redirect(url_for("hasil_generate", ujian_id=0, judul=judul, mapel=mapel, kelas=kelas, semester=semester, jumlah_soal=jumlah_soal, soal_output=soal_output))

    return render_template("buatsoal.html", soal_output=soal_output)

    
        # ... (lanjutkan redirect)

@app.route("/hasil_generate", methods=["GET"])
def hasil_generate():
    if "user" not in session or session["user"]["role"] != "pengajar":
        return redirect(url_for("login"))
    
    ujian_id = request.args.get("ujian_id")
    judul = request.args.get("judul")
    mapel = request.args.get("mapel")
    kelas = request.args.get("kelas")
    semester = request.args.get("semester")
    jumlah_soal = request.args.get("jumlah_soal")
    soal_output = request.args.get("soal_output")

    return render_template("hasil_generate.html", 
                           ujian_id=ujian_id, 
                           judul=judul, 
                           mapel=mapel, 
                           kelas=kelas, 
                           semester=semester,
                           jumlah_soal=jumlah_soal,
                           soal_output=soal_output)


@app.route('/save_generated', methods=['POST'])
def save_generated():
    if 'user' not in session or session['user']['role'] != 'pengajar':
        return redirect(url_for('login'))

    judul = request.form.get('judul')
    mapel = request.form.get('mapel')
    kelas = request.form.get('kelas')
    semester = request.form.get('semester')
    jumlah_soal = int(request.form.get('jumlah_soal', 0))
    soal_output = request.form.get('soal_output', '')

    # jika soal_output dikirim sebagai JSON string (escaped), coba unescape
    try:
        # jika dimulai dan berakhir dengan quotes, strip
        if (soal_output.startswith('"') and soal_output.endswith('"')) or (soal_output.startswith("'") and soal_output.endswith("'")):
            soal_output = soal_output[1:-1]
            soal_output = soal_output.encode('utf-8').decode('unicode_escape')
    except Exception:
        pass

    # Simpan ujian
    cursor.execute("""
        INSERT INTO ujian (judul, mapel, jumlah_soal, kelas, semester, dibuat_oleh)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (judul, mapel, jumlah_soal, kelas, semester, session['user']['id']))
    db.commit()
    ujian_id = cursor.lastrowid

    # Parsing sederhana: gunakan pola yang sama seperti sebelumnya
    pattern = r"Soal\s*\d+\s*:(.*?)Jawaban\s*:(.*?)Pembahasan\s*:(.*?)(?=Soal\s*\d+\s*:|$)"
    matches = re.findall(pattern, soal_output, re.S | re.I)

    for soal_text, jawaban_text, pembahasan_text in matches:
        cursor.execute("""
            INSERT INTO soal (ujian_id, soal, jawaban, pembahasan)
            VALUES (%s, %s, %s, %s)
        """, (ujian_id, soal_text.strip(), jawaban_text.strip(), pembahasan_text.strip()))
    db.commit()

    return ('', 200)  # return 200 OK untuk AJAX


@app.route('/report_soal', methods=['POST'])
def report_soal():
    # menerima laporan koreksi soal dari murid
    try:
        data = request.get_json() or {}
        ujian_id = data.get('ujian_id')
        soal_no = data.get('soal_no')
        komentar = data.get('komentar')
        murid_id = session['user']['id'] if 'user' in session else None

        report = {
            'ujian_id': ujian_id,
            'soal_no': soal_no,
            'komentar': komentar,
            'murid_id': murid_id
        }

        reports_file = os.path.join(os.path.dirname(__file__), 'reports.json')
        if os.path.exists(reports_file):
            with open(reports_file, 'r', encoding='utf-8') as f:
                existing = json.load(f)
        else:
            existing = []

        existing.append(report)
        with open(reports_file, 'w', encoding='utf-8') as f:
            json.dump(existing, f, ensure_ascii=False, indent=2)

        return ('', 204)
    except Exception as e:
        print('Error report_soal:', e)
        return ("Error", 500)


@app.route('/verify_face', methods=['POST'])
def verify_face():
    # Endpoint sederhana: menerima data:image/... dan kembalikan verified True
    try:
        data = request.get_json() or {}
        image_data = data.get('image')  # dataURL
        if not image_data:
            return {'verified': False}, 400

        # Simpan bukti verifikasi untuk audit (opsional)
        header, encoded = image_data.split(',', 1)
        img_bytes = base64.b64decode(encoded)
        fname = os.path.join(os.path.dirname(__file__), f"verify_{int(__import__('time').time())}.jpg")
        with open(fname, 'wb') as f:
            f.write(img_bytes)

        # Placeholder verification: terima selalu True (di masa depan, panggil face-recognition)
        return {'verified': True}
    except Exception as e:
        print('verify_face error:', e)
        return {'verified': False}, 500


# ---------------- EXPORT ----------------
@app.route("/export_word", methods=["POST"])
def export_word():
    soal_output = request.form.get("soal_output")
    doc = Document()
    doc.add_heading("Soal Ujian Nasional (Generated AI)", 0)
    doc.add_paragraph(soal_output)
    file_path = "soal_output.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)

@app.route("/export_pdf", methods=["POST"])
def export_pdf():
    soal_output = request.form.get("soal_output")
    file_path = "soal_output.pdf"
    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Soal Ujian Nasional (Generated AI)", styles['Title']),
        Paragraph(soal_output.replace("\n", "<br/>"), styles['Normal'])
    ]
    doc.build(story)
    return send_file(file_path, as_attachment=True)


if __name__== "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)